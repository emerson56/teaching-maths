from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent

NAVY = "17324D"
BLUE = "2E75B6"
LIGHT_BLUE = "DCEAF7"
PALE_BLUE = "EFF6FC"
GOLD = "D49B28"
LIGHT_GOLD = "FFF3D6"
GRAY = "5B6570"
LIGHT_GRAY = "EEF1F4"
MID_GRAY = "B9C1C9"
BLACK = "111111"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_border(cell, color=MID_GRAY, size="8", style="single"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), style)
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_in):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=0):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(total * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=11, bold=False, color=BLACK, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)
    sec.header_distance = Inches(0.25)
    sec.footer_distance = Inches(0.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in (
        ("Heading 1", 18, NAVY, 10, 6),
        ("Heading 2", 14, BLUE, 8, 4),
        ("Heading 3", 12, NAVY, 6, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_footer(doc, label):
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_font(run, size=8, color=GRAY)


def add_title(doc, title, subtitle=None, student_fields=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_font(r, size=24, bold=True, color=NAVY, name="Aptos Display")
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(8)
        r2 = p2.add_run(subtitle)
        set_font(r2, size=12, color=GRAY)
    if student_fields:
        table = doc.add_table(rows=1, cols=2)
        set_table_geometry(table, [3.8, 3.4])
        for cell, text in zip(table.rows[0].cells, ("Name: __________________________", "Date: __________________")):
            cell_border(cell, color=WHITE, size="0")
            p3 = cell.paragraphs[0]
            r3 = p3.add_run(text)
            set_font(r3, size=10, bold=True, color=GRAY)


def add_page_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text.upper())
    set_font(r, size=9, bold=True, color=BLUE)


def add_instruction(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [7.2])
    cell = table.cell(0, 0)
    shade(cell, PALE_BLUE)
    cell_border(cell, color=LIGHT_BLUE, size="8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, size=10.5, bold=True, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_kid_goal(doc, goal, tip=None):
    table = doc.add_table(rows=1, cols=2 if tip else 1)
    widths = [4.55, 2.65] if tip else [7.2]
    set_table_geometry(table, widths)
    cells = table.rows[0].cells
    shade(cells[0], LIGHT_BLUE)
    cell_border(cells[0], color=BLUE, size="9")
    p = cells[0].paragraphs[0]
    r1 = p.add_run("YOUR MISSION  ")
    set_font(r1, size=9, bold=True, color=BLUE)
    r2 = p.add_run(goal)
    set_font(r2, size=10.5, bold=True, color=NAVY)
    if tip:
        shade(cells[1], LIGHT_GOLD)
        cell_border(cells[1], color=GOLD, size="9")
        p2 = cells[1].paragraphs[0]
        r3 = p2.add_run("SMART MOVE  ")
        set_font(r3, size=9, bold=True, color=GOLD)
        r4 = p2.add_run(tip)
        set_font(r4, size=9.5, bold=True, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_lines(doc, count=3, label=None):
    if label:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label)
        set_font(r, size=10.5, bold=True, color=NAVY)
    for _ in range(count):
        p = doc.add_paragraph("________________________________________________________________________________")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            set_font(run, size=9, color=MID_GRAY)


def page_break(doc):
    doc.add_page_break()


def add_prompt(doc, number, text, space=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{number}. ")
    set_font(r, size=11, bold=True, color=BLUE)
    r2 = p.add_run(text)
    set_font(r2, size=11, color=BLACK)
    add_lines(doc, count=space)


def build_student():
    doc = Document()
    style_document(doc)
    add_footer(doc, "Grade 3 Summer Math | Week 1, Day 1")

    add_title(doc, "Place Value Explorer", "Your first number mission", student_fields=True)
    add_page_label(doc, "Page 1 | Warm-Up")
    add_kid_goal(
        doc,
        "Show what you already know about big numbers.",
        "A drawing or good guess is useful.",
    )
    add_instruction(doc, "Try these on your own. You are not expected to know everything yet.")
    add_prompt(doc, 1, "Read each number aloud: 4,706     3,090     8,215", 1)
    add_prompt(doc, 2, "Write 6,384 in expanded form.", 1)
    add_prompt(doc, 3, "Write 4,099 and 4,900 using words.", 2)
    add_prompt(doc, 4, "Place <, >, or = in the box:     4,099  [     ]  4,900", 1)
    add_prompt(doc, 5, "Round 647 to the nearest 10. Explain with a number line.", 2)
    add_prompt(doc, 6, "Round 647 to the nearest 100. Explain with a number line.", 2)

    page_break(doc)
    add_title(doc, "Mental Math Moves", "Find a clever route through each number.", student_fields=True)
    add_page_label(doc, "Page 2 | Number Moves")
    add_kid_goal(doc, "Solve in your head, then leave a trail that shows your idea.", "Friendly numbers can help.")
    add_instruction(doc, "Use words, jumps, or quick sketches. You do not need the vertical method.")
    for n, text in enumerate(
        (
            "399 + 58",
            "1,000 - 250",
            "75 + 68",
            "500 - 198",
            "What is 1 more, 10 more, and 100 more than 3,794?",
        ),
        1,
    ):
        add_prompt(doc, n, text, 2 if n < 5 else 3)
    add_lines(doc, 2, "Which problem felt easiest? Why?")

    page_break(doc)
    add_title(doc, "Build It! Place-Value Mat", "Turn a number into pieces, digits, and words.", student_fields=True)
    add_page_label(doc, "Page 3 | Reusable Math Tool")
    add_kid_goal(doc, "Put each value piece in its matching column.", "Build first. Add digits second.")
    add_instruction(doc, "Use the large boxes for pieces or drawings. Use the small boxes for one digit in each place.")
    table = doc.add_table(rows=3, cols=4)
    set_table_geometry(table, [1.8, 1.8, 1.8, 1.8])
    headers = ("Thousands", "Hundreds", "Tens", "Ones")
    for i, text in enumerate(headers):
        c = table.cell(0, i)
        shade(c, NAVY)
        cell_border(c, color=NAVY, size="10")
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, size=12, bold=True, color=WHITE)
        c2 = table.cell(1, i)
        shade(c2, "FFFFFF")
        cell_border(c2, color=BLUE, size="12")
        c2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        table.rows[1].height = Inches(2.05)
        table.rows[1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        c3 = table.cell(2, i)
        shade(c3, PALE_BLUE)
        cell_border(c3, color=BLUE, size="10")
        c3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        table.rows[2].height = Inches(0.55)
        table.rows[2].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        p3 = c3.paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run("digit: ______")
        set_font(r3, size=10.5, bold=True, color=NAVY)
    doc.add_paragraph()
    add_lines(doc, 1, "My number in standard form:")
    add_lines(doc, 1, "My number in expanded form:")
    add_lines(doc, 1, "My number in words:")

    page_break(doc)
    add_title(doc, "Build It, Say It, Write It", "Crack each number's place-value code.", student_fields=True)
    add_page_label(doc, "Page 4 | Guided Practice")
    add_kid_goal(doc, "Name every place and write the expanded form.", "Watch for empty places.")
    add_instruction(doc, "The first one is done as an example. Build the others on your mat before writing.")
    for number in ("3,245", "4,070", "6,308", "2,619", "5,040", "7,503"):
        table = doc.add_table(rows=1, cols=2)
        set_table_geometry(table, [1.25, 5.95])
        c1, c2 = table.rows[0].cells
        shade(c1, LIGHT_GOLD)
        shade(c2, "FFFFFF")
        cell_border(c1, color=GOLD)
        cell_border(c2, color=MID_GRAY)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(number)
        set_font(r1, size=15, bold=True, color=NAVY)
        p2 = c2.paragraphs[0]
        if number == "3,245":
            record = "3 thousands   2 hundreds   4 tens   5 ones"
            expanded = "Expanded: 3,000 + 200 + 40 + 5"
        else:
            record = "_____ thousands   _____ hundreds   _____ tens   _____ ones"
            expanded = "Expanded: __________________________________________________"
        r2 = p2.add_run(record)
        set_font(r2, size=10.5, color=BLACK)
        p3 = c2.add_paragraph(expanded)
        p3.paragraph_format.space_after = Pt(0)
        for run in p3.runs:
            set_font(run, size=10)

    page_break(doc)
    add_title(doc, "Number Name Match-Up", "One number can wear three different names.", student_fields=True)
    add_page_label(doc, "Page 5 | Independent Mission")
    add_kid_goal(doc, "Complete all three names for every number.", "Point across each row to check.")
    add_instruction(doc, "Fill every empty box. Then read each completed row aloud.")
    table = doc.add_table(rows=5, cols=3)
    set_table_geometry(table, [1.55, 2.55, 3.10])
    headers = ("Standard Form", "Expanded Form", "Word Form")
    for i, h in enumerate(headers):
        c = table.cell(0, i)
        shade(c, NAVY)
        cell_border(c, color=NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=10.5, bold=True, color=WHITE)
    rows = (
        ("2,619", "", ""),
        ("5,040", "", ""),
        ("", "7,000 + 500 + 3", ""),
        ("", "", "eight thousand four hundred sixty-two"),
    )
    for row_idx, row in enumerate(rows, 1):
        for col_idx, text in enumerate(row):
            c = table.cell(row_idx, col_idx)
            cell_border(c, color=MID_GRAY)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table.rows[row_idx].height = Inches(0.85)
            table.rows[row_idx].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_font(r, size=10.5, color=BLACK)
    add_lines(doc, 3, "Choose one number and explain the value of every digit:")

    page_break(doc)
    add_title(doc, "The Place-Value Trade Shop", "Change the pieces, but keep the same total.", student_fields=True)
    add_page_label(doc, "Page 6 | Trading Mission")
    add_kid_goal(doc, "Make a fair trade: 1 piece to the left equals 10 pieces to the right.", "Build the trade if unsure.")
    add_instruction(
        doc,
        "Decompose means break a number into place-value parts. The usual way uses one digit in each place.",
    )
    prompts = (
        "1 thousand = ______ hundreds",
        "1 hundred = ______ tens",
        "1 ten = ______ ones",
        "Example: 3,245 = 3 thousands + 2 hundreds + 4 tens + 5 ones",
        "3,245 = 2 thousands + ______ hundreds + 4 tens + 5 ones",
        "4,070 = 3 thousands + ______ hundreds + 7 tens",
    )
    for i, prompt in enumerate(prompts, 1):
        add_prompt(doc, i, prompt, 0 if i == 4 else 1)
    add_lines(doc, 1, "Break 8,462 apart in the usual way (8 thousands, 4 hundreds, ...):")
    add_lines(doc, 1, "Trade one thousand for 10 hundreds. Write the new way:")
    add_lines(doc, 1, "Why is the total still 8,462?")

    page_break(doc)
    add_title(doc, "The 14-Piece Puzzle", "How many different numbers can you unlock?", student_fields=True)
    add_page_label(doc, "Page 7 | Puzzle Mission")
    add_kid_goal(doc, "Use exactly 14 pieces to make a four-digit number.", "Keep each column from 0 to 9.")
    add_instruction(
        doc,
        "Use 0 to 9 pieces in each column. Add the four piece counts to get 14. The first row is an example.",
    )
    table = doc.add_table(rows=5, cols=6)
    set_table_geometry(table, [1.2, 1.2, 1.2, 1.2, 1.2, 1.2])
    headers = ("Thousands", "Hundreds", "Tens", "Ones", "Pieces", "Number")
    for i, h in enumerate(headers):
        c = table.cell(0, i)
        shade(c, NAVY)
        cell_border(c, color=NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=9, bold=True, color=WHITE)
    for row_idx in range(1, 5):
        table.rows[row_idx].height = Inches(0.52)
        table.rows[row_idx].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for c in table.rows[row_idx].cells:
            cell_border(c, color=MID_GRAY)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    example = ("5", "4", "3", "2", "14", "5,432")
    for col_idx, text in enumerate(example):
        c = table.cell(1, col_idx)
        shade(c, LIGHT_GOLD)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, size=11, bold=True, color=NAVY)
    add_lines(doc, 2, "How do you know each answer uses exactly 14 pieces?")
    add_lines(doc, 2, "What is the greatest number you can make? What is the least four-digit number?")

    page_break(doc)
    add_title(doc, "Final Number Mission", "Show the most important idea from today.", student_fields=True)
    add_page_label(doc, "Page 8 | Show What You Know")
    add_kid_goal(doc, "Prove that a digit changes value when it moves places.", "An equation or drawing can help.")
    add_instruction(doc, "Use complete sentences, equations, or drawings.")
    add_prompt(doc, 1, "In 5,204, what is the value of the digit 5?", 2)
    add_prompt(doc, 2, "In 52, what is the value of the digit 5?", 2)
    add_prompt(doc, 3, "Why is the 5 in 5,204 worth 100 times the 5 in 52?", 4)
    add_prompt(doc, 4, "Write one question you still have, or one interesting thing you noticed.", 3)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2.4, 2.4, 2.4])
    for cell, text, fill in zip(
        table.rows[0].cells,
        ("I feel secure", "I need more practice", "I need help"),
        (LIGHT_BLUE, LIGHT_GOLD, LIGHT_GRAY),
    ):
        shade(cell, fill)
        cell_border(cell, color=MID_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("□ " + text)
        set_font(r, size=10, bold=True, color=NAVY)

    page_break(doc)
    add_title(doc, "Cut-Apart Digit Cards", "Print on cardstock. Cut along the grid lines.")
    add_page_label(doc, "Page 9 | Reusable Manipulative")
    table = doc.add_table(rows=4, cols=5)
    set_table_geometry(table, [1.44] * 5)
    values = list("0123456789") * 2
    for idx, cell in enumerate([c for row in table.rows for c in row.cells]):
        cell_border(cell, color=NAVY, size="14", style="dashed")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row = table.rows[idx // 5]
        row.height = Inches(1.75)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(values[idx])
        set_font(r, size=38, bold=True, color=NAVY, name="Aptos Display")

    page_break(doc)
    add_title(doc, "Cut-Apart Thousands and Hundreds", "Cut on the dashed lines. Keep the pieces in labeled bags.")
    add_page_label(doc, "Page 10 | Reusable Math Tools")
    table = doc.add_table(rows=6, cols=4)
    set_table_geometry(table, [1.8] * 4)
    labels = (
        ("1,000", "1,000", "1,000", "1,000"),
        ("1,000", "1,000", "1,000", "1,000"),
        ("1,000", "100", "100", "100"),
        ("100", "100", "100", "100"),
        ("100", "100", "100", "100"),
        ("100", "100", "100", "100"),
    )
    fills = (NAVY, NAVY, BLUE, BLUE, LIGHT_BLUE, LIGHT_BLUE)
    text_colors = (WHITE, WHITE, WHITE, WHITE, NAVY, NAVY)
    for row_idx, row in enumerate(labels):
        table.rows[row_idx].height = Inches(1.10)
        table.rows[row_idx].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for col_idx, text in enumerate(row):
            c = table.cell(row_idx, col_idx)
            shade(c, fills[row_idx])
            cell_border(c, color=NAVY, size="12", style="dashed")
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_font(r, size=22 if text == "1,000" else 20, bold=True, color=text_colors[row_idx])

    page_break(doc)
    add_title(doc, "Cut-Apart Tens and Ones", "Cut on the dashed lines. Now you have enough pieces to build every lesson number.")
    add_page_label(doc, "Page 11 | Reusable Math Tools")
    table = doc.add_table(rows=5, cols=4)
    set_table_geometry(table, [1.8] * 4)
    labels = (
        ("10", "10", "10", "10"),
        ("10", "10", "10", "10"),
        ("10", "10", "1", "1"),
        ("1", "1", "1", "1"),
        ("1", "1", "1", "1"),
    )
    for row_idx, row in enumerate(labels):
        table.rows[row_idx].height = Inches(1.35)
        table.rows[row_idx].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for col_idx, text in enumerate(row):
            c = table.cell(row_idx, col_idx)
            fill = LIGHT_GOLD if text == "10" else LIGHT_GRAY
            shade(c, fill)
            cell_border(c, color=NAVY, size="12", style="dashed")
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_font(r, size=22, bold=True, color=NAVY)

    path = OUT / "week-01-day-01-student-packet.docx"
    doc.save(path)
    return path


def add_answer(doc, prompt, answer, note=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(prompt)
    set_font(r, size=9.5, bold=True, color=NAVY)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.2)
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(answer)
    set_font(r2, size=9.2, color=BLACK)
    if note:
        r3 = p2.add_run("  Observe: " + note)
        set_font(r3, size=8.5, color=GRAY)


def add_teacher_note(doc, label, text, fill=PALE_BLUE, border=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [7.2])
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell_border(cell, color=border, size="8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label + ": ")
    set_font(r1, size=9.2, bold=True, color=NAVY)
    r2 = p.add_run(text)
    set_font(r2, size=9.2, color=BLACK)


def add_teacher_bullets(doc, items, size=9.3):
    for text in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        set_font(r, size=size, color=BLACK)


def add_script(doc, teacher_says, purpose=None, listen_for=None, respond=None):
    rows = [("Teacher says", teacher_says, LIGHT_BLUE)]
    if purpose:
        rows.append(("Why", purpose, PALE_BLUE))
    if listen_for:
        rows.append(("Listen for", listen_for, LIGHT_GOLD))
    if respond:
        rows.append(("Then do", respond, LIGHT_GRAY))
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [1.25, 5.95])
    for row_idx, (label, text, fill) in enumerate(rows):
        left, right = table.rows[row_idx].cells
        shade(left, fill)
        cell_border(left, color=MID_GRAY)
        cell_border(right, color=MID_GRAY)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p1 = left.paragraphs[0]
        r1 = p1.add_run(label)
        set_font(r1, size=8.4, bold=True, color=NAVY)
        p2 = right.paragraphs[0]
        r2 = p2.add_run(text)
        set_font(r2, size=8.7, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_phase_header(doc, phase, minutes, student_pages, objective):
    table = doc.add_table(rows=2, cols=3)
    set_table_geometry(table, [3.75, 1.25, 2.20])
    values = (
        (phase, minutes, student_pages),
        (objective, "TIME", "MATERIAL"),
    )
    for row_idx, row in enumerate(values):
        for col_idx, text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell_border(cell, color=NAVY if row_idx == 0 else MID_GRAY)
            shade(cell, NAVY if row_idx == 0 else PALE_BLUE)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_font(
                r,
                size=12 if row_idx == 0 and col_idx == 0 else 9,
                bold=True,
                color=WHITE if row_idx == 0 else NAVY,
            )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_teacher_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_font(r, size=19, bold=True, color=NAVY, name="Aptos Display")
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        p2.paragraph_format.keep_with_next = True
        r2 = p2.add_run(subtitle)
        set_font(r2, size=10, color=GRAY)


def build_teacher():
    doc = Document()
    style_document(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(9.3)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0
    for style_name, size, before, after in (
        ("Heading 1", 14.5, 7, 3),
        ("Heading 2", 12, 6, 3),
        ("Heading 3", 10.5, 4, 2),
    ):
        style = doc.styles[style_name]
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    add_footer(doc, "Teacher Guide | Grade 3 Summer Math | Week 1, Day 1")
    add_teacher_title(doc, "Teaching Manual: Place Value Explorer", "Week 1, Day 1 | Scripted validation lesson | 60-75 minutes")
    add_page_label(doc, "Start Here")
    add_instruction(doc, "Lesson goal: the student understands that a digit's value is determined by its position in a base-ten system.")

    doc.add_heading("What the Student Should Learn", level=1)
    add_teacher_bullets(doc, (
        "Read and describe four-digit numbers using thousands, hundreds, tens, and ones.",
        "Connect a concrete model, a place-value chart, standard form, expanded form, and word form.",
        "Understand that 10 units in one place can be traded for 1 unit in the place to the left.",
        "Explain that moving a digit one place left makes its value 10 times as great.",
    ))

    doc.add_heading("Success Criteria", level=1)
    add_teacher_bullets(doc, (
        "The student says, for example, 'The 2 in 3,245 is worth 200,' not merely 'It is in the hundreds place.'",
        "The student treats zero as a placeholder in numbers such as 4,070 and 6,308.",
        "The student can trade 1 thousand for 10 hundreds without believing the total changed.",
        "The student can represent one number in at least three forms.",
    ))

    add_teacher_note(
        doc,
        "Important",
        "Do not read every script word-for-word if the child is already reasoning well. The script gives the intended language and decision points. Pause often and let the student do the mathematical talking.",
        fill=LIGHT_GOLD,
        border=GOLD,
    )

    doc.add_heading("Lesson at a Glance", level=1)
    table = doc.add_table(rows=6, cols=3)
    set_table_geometry(table, [1.05, 4.55, 1.60])
    rows = (
        ("10 min", "Diagnostic: number sense and mental math", "Student pp. 1-2"),
        ("7 min", "Number talk: make 1,000", "Loose pieces"),
        ("15 min", "Model 3,245 and connect representations", "Mat + pieces"),
        ("12 min", "Zero placeholders and trading", "Mat + digit cards"),
        ("15 min", "Guided and independent practice", "Student pp. 4-6"),
        ("10 min", "Challenge and exit ticket", "Student pp. 7-8"),
    )
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = table.cell(i, j)
            cell_border(cell, color=MID_GRAY)
            if j == 0:
                shade(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            set_font(r, size=9.2, bold=(j == 0), color=NAVY if j == 0 else BLACK)

    add_teacher_title(doc, "The Mathematics and Teaching Theory", "What the teacher needs to understand before teaching")
    add_page_label(doc, "Theory")
    doc.add_heading("1. Place Value Is Multiplicative", level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        "Our number system is base ten. Ten ones compose one ten, ten tens compose one hundred, and ten hundreds compose one thousand. "
        "A digit therefore has two meanings: its face value and the value created by its position. In 3,245, the digit 2 has face value 2 but place value 200."
    )
    set_font(r, size=10.5)
    add_teacher_note(doc, "Teacher implication", "Always ask both 'What digit do you see?' and 'What is that digit worth?'")

    doc.add_heading("2. Concrete -> Pictorial -> Abstract", level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        "The student first handles or sees quantities, then organizes them on a place-value mat, and only then records symbols. "
        "The materials are not decoration: they make the ten-to-one relationship visible."
    )
    set_font(r, size=10.5)
    add_teacher_bullets(doc, (
        "Concrete: blocks, disks, or cut-apart value pieces.",
        "Pictorial: drawings or labeled columns on the place-value mat.",
        "Abstract: 3,245; 3,000 + 200 + 40 + 5; number words.",
    ))

    doc.add_heading("3. Variation Reveals Understanding", level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        "The sequence 3,245 -> 4,070 -> 6,308 is deliberate. The first number has all four places occupied. "
        "The next two contain internal zeros, which reveal whether the student understands placeholders or simply reads digits from left to right."
    )
    set_font(r, size=10.5)

    doc.add_heading("4. Explanation Is Part of the Mathematics", level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        "A correct answer does not by itself prove secure understanding. Ask the child to point, build, compare, and explain. "
        "Avoid supplying vocabulary too early; first listen to the child's idea, then refine the language."
    )
    set_font(r, size=10.5)

    add_teacher_title(doc, "Materials: Preparation and Use", "Set up the table before the student arrives")
    add_page_label(doc, "Materials")
    doc.add_heading("Print and Prepare", level=1)
    add_teacher_bullets(doc, (
        "Print student pages 1-8 single-sided.",
        "Print page 9 digit cards and pages 10-11 place-value pieces on cardstock, then cut on dashed lines.",
        "Put student page 3, the place-value mat, in a sheet protector if possible.",
        "Prepare a pencil, eraser, dry-erase marker, and two blank sheets for teacher modeling.",
        "Keep the answer pages out of the student's view.",
    ))

    doc.add_heading("What Each Material Does", level=1)
    table = doc.add_table(rows=5, cols=3)
    set_table_geometry(table, [1.75, 2.25, 3.20])
    material_rows = (
        ("Place-value mat", "Organizes units by position", "Put one type of piece in each labeled column. Never mix hundreds into the tens column."),
        ("Value pieces", "Represent actual quantities", "A card marked 100 means one hundred, not one object to count as 1."),
        ("Digit cards", "Represent standard notation", "Place one digit below each column only after the quantity has been built."),
        ("Student packet", "Records thinking", "Use pp. 1-2 diagnostically; pp. 4-6 for practice; pp. 7-8 for transfer and assessment."),
        ("Blocks/disks", "Optional stronger concrete model", "Use in place of printed pieces. Insist on a consistent value for every shape or color."),
    )
    for row_idx, row in enumerate(material_rows):
        for col_idx, text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell_border(cell, color=MID_GRAY)
            if col_idx == 0:
                shade(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            set_font(r, size=9, bold=(col_idx == 0), color=NAVY if col_idx == 0 else BLACK)

    doc.add_heading("Table Layout", level=1)
    add_teacher_bullets(doc, (
        "Center: place-value mat.",
        "Above the mat: piles of thousands, hundreds, tens, and ones, clearly separated.",
        "Below the mat: digit cards 0-9.",
        "Student's dominant-hand side: packet and pencil.",
        "Teacher's side: observation notes. Do not write on the student's paper during the diagnostic.",
    ))
    add_teacher_note(
        doc,
        "Material rule",
        "When the student builds a number, require pieces to be placed in the matching column. When trading, physically remove one piece and replace it with 10 equal-value pieces before writing an equation.",
        fill=LIGHT_GOLD,
        border=GOLD,
    )

    add_teacher_title(doc, "Phase 1: Welcome and Diagnostic", "Observe before teaching")
    add_page_label(doc, "Student Pages 1-2")
    add_phase_header(doc, "Diagnostic", "8-12 min", "Packet pp. 1-2", "Find out what the student already understands.")
    doc.add_heading("Set the Tone", level=1)
    add_script(
        doc,
        '"Today I want to see how you think about numbers. Some questions may be easy and some may be unfamiliar. This is not a grade. If you are unsure, show any idea you have."',
        purpose="This reduces performance pressure and makes the diagnostic more truthful.",
        listen_for="A willingness to attempt, draw, or explain.",
        respond="Give neutral encouragement: 'Thank you for showing your thinking.' Do not say whether an answer is right yet.",
    )

    doc.add_heading("How to Administer Page 1", level=1)
    add_teacher_bullets(doc, (
        "Ask the student to read 4,706; 3,090; and 8,215 aloud. Do not pronounce them first.",
        "For expanded form and comparison, say only: 'Show what you know.'",
        "For rounding, let the student draw any number line they understand.",
        "If the student asks what a word means, clarify ordinary language but do not teach the mathematics.",
    ))

    doc.add_heading("How to Administer Page 2", level=1)
    add_script(
        doc,
        '"Please solve these in your head as much as you can. Write a few marks or words so I can see the route your mind took. You do not need to use the standard vertical method."',
        purpose="The goal is to reveal number relationships and flexibility, not speed.",
        listen_for="Compensation, making friendly numbers, decomposing tens and ones, or counting by ones.",
        respond="After each problem ask only, 'How did you get that?' Record the strategy without correcting it.",
    )

    doc.add_heading("Diagnostic Observation Codes", level=1)
    table = doc.add_table(rows=4, cols=3)
    set_table_geometry(table, [1.05, 2.65, 3.50])
    codes = (
        ("S", "Secure", "Accurate and explains using place value or an efficient relationship."),
        ("D", "Developing", "Mostly accurate but relies on prompts, counting, or imprecise language."),
        ("F", "Fragile", "Confuses places, ignores zeros, or cannot connect a digit to its value."),
        ("N/O", "Not observed", "No attempt or insufficient evidence; revisit later without treating it as failure."),
    )
    for i, row in enumerate(codes):
        for j, text in enumerate(row):
            cell = table.cell(i, j)
            cell_border(cell, color=MID_GRAY)
            if j == 0:
                shade(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            set_font(r, size=9, bold=(j == 0), color=NAVY if j == 0 else BLACK)

    add_teacher_title(doc, "Using the Diagnostic", "Decide how much support to give in the lesson")
    add_page_label(doc, "Teacher Decision")
    doc.add_heading("What the Responses Mean", level=1)
    table = doc.add_table(rows=5, cols=3)
    set_table_geometry(table, [2.15, 2.55, 2.50])
    diagnostic_rows = (
        ("Reads 3,090 as 'three thousand ninety'", "Understands zero placeholders in speech", "Continue to 4,070 and 6,308."),
        ("Reads it as 'three thousand nine hundred'", "Hundreds and tens are being confused", "Return to the mat and point to every column."),
        ("Writes 6,384 as 6,000 + 300 + 80 + 4", "Connects digit position to value", "Ask for a nonstandard decomposition later."),
        ("Solves 399 + 58 as 400 + 58 - 1", "Uses compensation flexibly", "Invite multiple strategies, not extra routine work."),
        ("Counts 399 + 58 by ones", "Calculation is possible but inefficient", "Do not reteach now; note for Day 4 mental-math work."),
    )
    for row_idx, row in enumerate(diagnostic_rows):
        for col_idx, text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell_border(cell, color=MID_GRAY)
            if col_idx == 0:
                shade(cell, PALE_BLUE)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            set_font(r, size=8.8, bold=(col_idx == 0), color=NAVY if col_idx == 0 else BLACK)

    doc.add_heading("Choose the Lesson Path", level=1)
    add_teacher_note(
        doc,
        "If secure",
        "Move briskly through the first model. Ask for alternate decompositions and justifications.",
    )
    add_teacher_note(
        doc,
        "If developing",
        "Teach the lesson as scripted. Keep pieces on the mat while recording symbols.",
        fill=LIGHT_GOLD,
        border=GOLD,
    )
    add_teacher_note(
        doc,
        "If fragile",
        "Reduce the first number to 245. Establish hundreds, tens, and ones before adding thousands. Skip the hardest challenge if needed.",
        fill=LIGHT_GRAY,
        border=MID_GRAY,
    )
    doc.add_heading("Transition to Teaching", level=1)
    add_script(
        doc,
        '"Thank you. I learned something about how you see numbers. Now we are going to investigate why the same digit can be worth very different amounts depending on where it is placed."',
        purpose="Clearly closes assessment mode and opens collaborative learning mode.",
    )

    add_teacher_title(doc, "Phase 2: Number Talk", "Compose 1,000 in many ways")
    add_page_label(doc, "Oral Reasoning")
    add_phase_header(doc, "Number Talk", "5-8 min", "Loose value pieces", "Establish 10-to-1 trading before four-digit modeling.")
    doc.add_heading("Launch", level=1)
    add_script(
        doc,
        '"Here is one thousand. How else could we make exactly the same amount using hundreds, tens, or ones?"',
        purpose="Elicit equivalence: the appearance and number of pieces can change while total value stays fixed.",
        listen_for="'10 hundreds,' '100 tens,' or mixed decompositions.",
        respond="After each answer ask, 'How could we prove it?' Build or sketch the proposed trade.",
    )

    doc.add_heading("Use the Materials", level=1)
    add_teacher_bullets(doc, (
        "Place one 1,000 piece in the thousands column.",
        "Remove it and place ten 100 pieces in the hundreds column. Count by hundreds: 100, 200, ... 1,000.",
        "If pieces are limited, draw ten labeled hundred boxes rather than pretending fewer pieces are enough.",
        "Record beside the mat: 1,000 = 10 hundreds.",
        "Then ask how 10 hundreds could become tens. You may record 1,000 = 100 tens without laying out all 100 pieces.",
    ))

    doc.add_heading("Productive Follow-Ups", level=1)
    add_script(doc, '"Could 9 hundreds equal 1,000? What is missing?"', listen_for="100 more, or 10 tens.")
    add_script(doc, '"If I use more pieces, does that always mean I have a greater number?"', listen_for="No; piece value matters.")
    add_script(doc, '"What stayed the same when we traded?"', listen_for="The total value.")

    add_teacher_note(
        doc,
        "Do not say",
        "'We add a zero' or 'Move it over.' Those shortcuts hide the quantity relationship. Say, 'Ten of this unit compose one of the next unit.'",
        fill=LIGHT_GOLD,
        border=GOLD,
    )

    add_teacher_title(doc, "Phase 3: Model 3,245", "Connect quantity, position, and notation")
    add_page_label(doc, "Student Page 3")
    add_phase_header(doc, "Explicit Model", "12-15 min", "Mat, pieces, digit cards", "Build one complete four-digit example.")
    doc.add_heading("Step 1: Build the Quantity", level=1)
    add_script(
        doc,
        '"We are going to build three thousand two hundred forty-five. Before touching the pieces, what units do you expect we will need?"',
        purpose="Prediction activates the structure before manipulation.",
        listen_for="Thousands, hundreds, tens, and ones.",
        respond="Have the student place 3 thousand pieces, 2 hundred pieces, 4 ten pieces, and 5 one pieces in the matching columns.",
    )
    add_script(
        doc,
        '"Point to the digit 2. What digit is it? What is it worth in this number? How can the pieces prove that?"',
        listen_for="'It is the digit 2, but it is worth 200 because there are two hundreds.'",
        respond="If the child says only 'hundreds,' prompt: 'How many hundreds? What total value?'",
    )

    doc.add_heading("Step 2: Add Digit Cards", level=1)
    add_teacher_bullets(doc, (
        "Place digit card 3 below the thousands column, 2 below hundreds, 4 below tens, and 5 below ones.",
        "Slide the cards together to form 3,245.",
        "Point from each card to its corresponding pieces.",
        "Do not place the digit cards first; quantity should lead notation in this first model.",
    ))

    doc.add_heading("Step 3: Record Four Forms", level=1)
    add_script(
        doc,
        '"We built one number. Now we will name the same number in several languages."',
        purpose="The child should see forms as equivalent representations, not separate facts.",
        respond="On student page 3 record: 3,245; 3,000 + 200 + 40 + 5; three thousand two hundred forty-five; 3 thousands, 2 hundreds, 4 tens, 5 ones.",
    )
    add_teacher_note(
        doc,
        "Check",
        "Cover the pieces and ask the student to rebuild from 3,245. Then cover the numeral and ask the student to write it from the pieces.",
    )

    add_teacher_title(doc, "Phase 4: Zeros and Placeholders", "Teach 4,070 and 6,308")
    add_page_label(doc, "Variation")
    add_phase_header(doc, "Placeholder Models", "8-10 min", "Mat + digit cards", "Reveal whether zero is understood as an occupied place with no units.")
    doc.add_heading("Model 4,070", level=1)
    add_script(
        doc,
        '"Build four thousand seventy. Which columns will contain pieces? Which column must remain empty?"',
        listen_for="4 thousands, 0 hundreds, 7 tens, 0 ones.",
        respond="Leave the hundreds and ones columns visibly empty. Place zero digit cards below those columns.",
    )
    add_script(
        doc,
        '"Why can we not write 47 for this model?"',
        purpose="Make the placeholder function explicit.",
        listen_for="'47 would mean 4 tens and 7 ones. The zeros hold the thousands and tens in their correct positions.'",
        respond="Compare 4,070 and 47 side-by-side with the mat.",
    )

    doc.add_heading("Model 6,308", level=1)
    add_script(
        doc,
        '"Now build six thousand three hundred eight. Tell me what belongs in every column, including the empty one."',
        listen_for="6 thousands, 3 hundreds, 0 tens, 8 ones.",
        respond="Ask the student to place the zero card in the tens position and read the completed number aloud.",
    )

    doc.add_heading("Common Misconceptions", level=1)
    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [3.10, 4.10])
    misconceptions = (
        ("Student reads 4,070 as 'four thousand seven hundred.'", "Point column by column: 'How many hundreds are shown? How many tens?' Rebuild rather than merely correcting the reading."),
        ("Student omits zero and writes 470.", "Place 4 in the thousands column. Ask, 'If we remove this zero, where does the 4 move, and what would it be worth?'"),
        ("Student says zero has no purpose.", "Agree that it represents no pieces, then add: 'It still holds a place so the other digits keep their values.'"),
        ("Student counts printed pieces as single objects.", "Ask the child to read the label on each piece and count by its value: 100, 200, 300."),
    )
    for row_idx, row in enumerate(misconceptions):
        for col_idx, text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell_border(cell, color=MID_GRAY)
            if col_idx == 0:
                shade(cell, LIGHT_GOLD)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            set_font(r, size=9, bold=(col_idx == 0), color=NAVY if col_idx == 0 else BLACK)

    add_teacher_title(doc, "Phase 5: Trading Without Changing Value", "Make regrouping conceptually visible")
    add_page_label(doc, "Student Page 6")
    add_phase_header(doc, "Trading", "8-10 min", "Mat + pieces", "Show that a number can be decomposed in more than one equivalent way.")
    doc.add_heading("Demonstrate the Trade", level=1)
    add_script(
        doc,
        '"Look at 3,245. I am going to remove one thousand piece. Have I changed the number yet?"',
        listen_for="'Yes, it is now 1,000 less.'",
        respond="Confirm: 'Right now it has changed. To preserve the value, what must replace that thousand?'",
    )
    add_script(
        doc,
        '"Replace the one thousand with ten hundreds. Count the hundreds now. Did the total number change? Prove it."',
        purpose="Regrouping is an exchange of equal value, not a digit trick.",
        listen_for="12 hundreds total; 2,000 + 1,200 + 40 + 5 still equals 3,245.",
        respond="Record both decompositions directly under one another.",
    )

    doc.add_heading("Write the Equivalence", level=1)
    add_teacher_bullets(doc, (
        "Usual form: 3 thousands + 2 hundreds + 4 tens + 5 ones.",
        "After trading: 2 thousands + 12 hundreds + 4 tens + 5 ones.",
        "Expanded values: 3,000 + 200 + 40 + 5 = 2,000 + 1,200 + 40 + 5.",
        "Ask the student to circle the part that changed and underline the total that stayed equal.",
    ))

    doc.add_heading("Guided Example: 4,070", level=1)
    add_script(
        doc,
        '"Trade one thousand in 4,070 for hundreds. What is the new decomposition?"',
        listen_for="3 thousands + 10 hundreds + 7 tens.",
        respond="If the student writes 3,107, return to units: 3 thousands, 10 hundreds, 7 tens is not read by concatenating the counts.",
    )

    add_teacher_note(
        doc,
        "Why this matters later",
        "This physical exchange is the foundation for regrouping in addition and subtraction. A child who understands the exchange will be less dependent on memorized carrying and borrowing rules.",
        fill=LIGHT_GOLD,
        border=GOLD,
    )

    add_teacher_title(doc, "Phase 6: Guided and Independent Practice", "Move responsibility gradually to the student")
    add_page_label(doc, "Student Pages 4-6")
    add_phase_header(doc, "Practice", "12-15 min", "Packet + mat", "Shift from teacher model to independent representation.")
    doc.add_heading("Use an I Do -> We Do -> You Do Sequence", level=1)
    table = doc.add_table(rows=3, cols=3)
    set_table_geometry(table, [1.25, 2.55, 3.40])
    practice_rows = (
        ("I do", "3,245", "Teacher models briefly while student explains each placement."),
        ("We do", "4,070 and 6,308", "Student places pieces; teacher asks questions and records only after agreement."),
        ("You do", "2,619; 5,040; 7,503", "Student builds and records. Teacher observes and prompts minimally."),
    )
    for row_idx, row in enumerate(practice_rows):
        for col_idx, text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell_border(cell, color=MID_GRAY)
            if col_idx == 0:
                shade(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            set_font(r, size=9, bold=(col_idx == 0), color=NAVY if col_idx == 0 else BLACK)

    doc.add_heading("Prompt Ladder", level=1)
    add_teacher_note(doc, "First prompt", "What do you notice? What could you try?")
    add_teacher_note(doc, "Second prompt", "Which place should you examine first?", fill=PALE_BLUE)
    add_teacher_note(doc, "Third prompt", "Point to the column for hundreds. What belongs there?", fill=LIGHT_GOLD, border=GOLD)
    add_teacher_note(doc, "Last resort", "Model a smaller related number, then return to the original.", fill=LIGHT_GRAY, border=MID_GRAY)

    doc.add_heading("What to Say During Page 5", level=1)
    add_script(
        doc,
        '"These are not three different numbers. They are three names for the same number. After completing a row, point across it and prove that every form has the same value."',
        listen_for="Connections between each digit and its expanded value.",
        respond="For word form, have the student read aloud and point to each place. Do not penalize handwriting or spelling if the mathematical wording is clear.",
    )

    doc.add_heading("What to Say During Page 6", level=1)
    add_script(
        doc,
        '"Your second decomposition must look different but have exactly the same total. Show the trade that makes it legal."',
        listen_for="A valid exchange of 1 unit for 10 units to the right.",
        respond="Reject unexplained digit rearrangement. Ask the student to prove equality with pieces or expanded values.",
    )

    add_teacher_title(doc, "Phase 7: Fourteen-Piece Challenge", "Competition habit: organize possibilities")
    add_page_label(doc, "Student Page 7")
    add_phase_header(doc, "Challenge", "7-10 min", "Packet p. 7 + pieces", "Search systematically and justify multiple solutions.")
    doc.add_heading("Launch Without Giving a Method", level=1)
    add_script(
        doc,
        '"A piece may be worth 1,000, 100, 10, or 1. Build a four-digit number using exactly 14 pieces. There is more than one answer. Start anywhere, but be ready to prove your count."',
        purpose="The problem develops organized search, place-value composition, and justification.",
        listen_for="The student checks that thousands + hundreds + tens + ones counts total 14.",
        respond="Wait at least two minutes before suggesting a strategy.",
    )

    doc.add_heading("If the Student Is Stuck", level=1)
    add_script(doc, '"A four-digit number needs at least how many thousand pieces?"', listen_for="At least 1.")
    add_script(doc, '"Suppose you choose 3 thousands. How many pieces remain for the other columns?"', listen_for="11.")
    add_script(doc, '"How could a table help you avoid repeating an answer?"', respond="Use the six-column table on page 7.")

    doc.add_heading("Extend the Reasoning", level=1)
    add_script(
        doc,
        '"What is the greatest ordinary four-digit number possible with 14 pieces? What is the least? Convince me no better answer exists."',
        listen_for="Greatest 9,500; least 1,049, assuming each column is written with a standard digit 0-9.",
        respond="Ask the student to explain why pieces should be moved toward the left for the greatest number and toward the right for the least.",
    )
    add_teacher_note(
        doc,
        "Competition habit",
        "Praise organization and proof, not just finding an answer: 'Your table makes it clear you did not repeat a case.'",
        fill=LIGHT_GOLD,
        border=GOLD,
    )

    add_teacher_title(doc, "Phase 8: Exit Ticket and Feedback", "Assess the central place-value idea")
    add_page_label(doc, "Student Page 8")
    add_phase_header(doc, "Exit Ticket", "5-7 min", "Packet p. 8", "Decide whether the student understands tenfold place-value relationships.")
    doc.add_heading("Give the Exit Ticket Independently", level=1)
    add_script(
        doc,
        '"Please answer these without my help. Use words, an equation, or a drawing to make your explanation convincing."',
        purpose="The teacher needs independent evidence after instruction.",
        respond="Do not prompt until the student has finished. Then ask one clarifying question if the writing is ambiguous.",
    )

    doc.add_heading("Expected Explanation", level=1)
    add_teacher_note(
        doc,
        "Strong answer",
        "The 5 in 5,204 is worth 5,000. The 5 in 52 is worth 50. The first 5 is two places farther left, and each move left multiplies value by 10. Therefore 5,000 is 100 times 50.",
    )
    doc.add_heading("Respond to the Student", level=1)
    table = doc.add_table(rows=3, cols=3)
    set_table_geometry(table, [1.25, 2.55, 3.40])
    feedback_rows = (
        ("Secure", "Correct value and tenfold explanation", 'Say: "You connected position to multiplication by 10. Tomorrow we can push that idea further."'),
        ("Developing", "Correct values but weak explanation", 'Say: "Your values are correct. Let us use the mat once more to show why the relationship is 100 times."'),
        ("Fragile", "Calls both digits 5 or confuses 500/5,000", 'Say: "The digit is the same, but its value changes. Tomorrow we will begin by rebuilding both numbers."'),
    )
    for row_idx, row in enumerate(feedback_rows):
        for col_idx, text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell_border(cell, color=MID_GRAY)
            if col_idx == 0:
                shade(cell, LIGHT_BLUE if row_idx == 0 else LIGHT_GOLD if row_idx == 1 else LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            set_font(r, size=8.8, bold=(col_idx == 0), color=NAVY if col_idx == 0 else BLACK)

    doc.add_heading("Close Positively", level=1)
    add_script(
        doc,
        '"Today you showed that a number can look different in pieces, words, and equations while keeping the same value. Tell me one idea from today that you could teach someone else."',
        purpose="Retrieval and student teaching strengthen memory while ending on competence.",
    )

    add_teacher_title(doc, "Answer Key and Observation Record", "Use after the student has completed the work")
    add_page_label(doc, "Reference")
    doc.add_heading("Diagnostic Answers", level=1)
    add_answer(doc, "Read numbers", "4,706: four thousand seven hundred six. 3,090: three thousand ninety. 8,215: eight thousand two hundred fifteen.")
    add_answer(doc, "Expanded form", "6,384 = 6,000 + 300 + 80 + 4.")
    add_answer(doc, "Comparison", "4,099 < 4,900.")
    add_answer(doc, "Rounding", "647 -> 650 to nearest 10; 647 -> 600 to nearest 100.")
    add_answer(doc, "Mental math", "399 + 58 = 457; 1,000 - 250 = 750; 75 + 68 = 143; 500 - 198 = 302.")
    add_answer(doc, "More than 3,794", "1 more: 3,795. 10 more: 3,804. 100 more: 3,894.")

    doc.add_heading("Practice Answers", level=1)
    table = doc.add_table(rows=6, cols=2)
    set_table_geometry(table, [1.35, 5.85])
    answers = (
        ("3,245", "3 thousands, 2 hundreds, 4 tens, 5 ones"),
        ("4,070", "4 thousands, 0 hundreds, 7 tens, 0 ones"),
        ("6,308", "6 thousands, 3 hundreds, 0 tens, 8 ones"),
        ("2,619", "2 thousands, 6 hundreds, 1 ten, 9 ones"),
        ("5,040", "5 thousands, 0 hundreds, 4 tens, 0 ones"),
        ("7,503", "7 thousands, 5 hundreds, 0 tens, 3 ones"),
    )
    for row_idx, (number, text) in enumerate(answers):
        for col_idx, value in enumerate((number, text)):
            cell = table.cell(row_idx, col_idx)
            cell_border(cell, color=MID_GRAY)
            if col_idx == 0:
                shade(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            r = p.add_run(value)
            set_font(r, size=9, bold=(col_idx == 0), color=NAVY if col_idx == 0 else BLACK)

    doc.add_heading("Trading and Challenge", level=1)
    add_answer(doc, "3,245 after one thousand trade", "2 thousands + 12 hundreds + 4 tens + 5 ones.")
    add_answer(doc, "4,070 after one thousand trade", "3 thousands + 10 hundreds + 7 tens.")
    add_answer(doc, "8,462 alternate example", "7 thousands + 14 hundreds + 6 tens + 2 ones.")
    add_answer(doc, "Fourteen-piece examples", "5,432 and 2,741 each use exactly 14 pieces.")
    add_answer(doc, "Challenge extremes", "Greatest: 9,500. Least four-digit number: 1,049.")

    add_teacher_title(doc, "Teacher Observation Record", "Complete immediately after the lesson")
    add_page_label(doc, "Planning for Day 2")
    doc.add_heading("Evidence Checklist", level=1)
    checks = (
        "Reads numbers containing internal zeros accurately.",
        "Names both the digit and its value.",
        "Matches pieces to the correct place-value column.",
        "Uses zero as a placeholder.",
        "Trades 1 unit for 10 units to the right without changing total value.",
        "Writes standard, expanded, word, and unit forms.",
        "Explains a tenfold relationship.",
        "Persists and organizes work during the challenge.",
    )
    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [3.6, 3.6])
    for idx, text in enumerate(checks):
        cell = table.cell(idx // 2, idx % 2)
        cell_border(cell, color=MID_GRAY)
        p = cell.paragraphs[0]
        r = p.add_run("□ " + text)
        set_font(r, size=9, color=BLACK)

    doc.add_heading("Record Specific Evidence", level=1)
    add_lines(doc, 3, "A strong statement or strategy the student used:")
    add_lines(doc, 3, "A misconception or hesitation to revisit:")
    add_lines(doc, 3, "Material that helped most:")

    doc.add_heading("Decision for Day 2", level=1)
    add_teacher_note(doc, "Secure", "Proceed to expanded form and comparing numbers; include extension questions.")
    add_teacher_note(doc, "Developing", "Proceed, but keep the mat and digit cards available.", fill=LIGHT_GOLD, border=GOLD)
    add_teacher_note(doc, "Fragile", "Begin Day 2 with a 10-minute small-number reteach using 245, 407, and 608.", fill=LIGHT_GRAY, border=MID_GRAY)

    path = OUT / "week-01-day-01-teacher-guide.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(build_student())
    print(build_teacher())
