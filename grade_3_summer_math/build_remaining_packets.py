from __future__ import annotations

import importlib.util
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
DAY1_BUILDER = ROOT / "week-01-place-value" / "materials" / "build_day01_packet.py"
spec = importlib.util.spec_from_file_location("day1_helpers", DAY1_BUILDER)
h = importlib.util.module_from_spec(spec)
spec.loader.exec_module(h)


DOMAIN = {
    1: {
        "name": "place value",
        "theory": "Base-ten understanding links each place to the place beside it by a factor of 10. Students should connect quantities, visual models, spoken language, and symbols instead of memorizing isolated rules.",
        "vocab": ["place value", "standard form", "expanded form", "compare", "round", "estimate"],
        "misconceptions": [
            "Reading internal zeros as if they were occupied places.",
            "Comparing numbers by the last digit instead of the greatest place.",
            "Rounding by a rhyme without locating benchmarks.",
        ],
        "smart": "Start with the greatest place and explain what stays the same.",
        "tool": "number_line",
    },
    2: {
        "name": "addition and subtraction",
        "theory": "Regrouping is an exchange of equal value: 10 ones compose 1 ten and 10 tens compose 1 hundred. Written algorithms should record exchanges the student can first model and explain.",
        "vocab": ["addend", "sum", "difference", "regroup", "inverse operation", "bar model"],
        "misconceptions": [
            "Treating regrouped digits as unexplained small numbers.",
            "Subtracting the smaller digit from the larger regardless of position.",
            "Choosing an operation from keywords instead of the story structure.",
        ],
        "smart": "Estimate first, then use place value to check every trade.",
        "tool": "place_grid",
    },
    3: {
        "name": "multiplication foundations",
        "theory": "Multiplication describes equal groups and arrays. Fluency grows from relationships among facts, while division reverses the structure by finding a group size or number of groups.",
        "vocab": ["equal groups", "factor", "product", "array", "row", "column", "fact family"],
        "misconceptions": [
            "Counting unequal groups as multiplication.",
            "Confusing the number of groups with the size of each group.",
            "Memorizing a product without being able to model it.",
        ],
        "smart": "Draw groups or an array before relying on memory.",
        "tool": "array_grid",
    },
    4: {
        "name": "multiplication and division",
        "theory": "Harder facts should be derived from known facts through decomposition, doubling, and the distributive property. Division is most secure when linked to an unknown multiplication factor and interpreted in context.",
        "vocab": ["factor", "product", "quotient", "divisor", "remainder", "unknown factor"],
        "misconceptions": [
            "Treating division only as sharing and missing grouping situations.",
            "Ignoring what a remainder means in the story.",
            "Using a fact trick without understanding why it works.",
        ],
        "smart": "Turn division into a missing multiplication fact.",
        "tool": "fact_grid",
    },
    5: {
        "name": "properties and algebraic thinking",
        "theory": "Properties describe why equivalent expressions keep the same value. Variables represent unknown quantities, and the equal sign states that two expressions have the same value.",
        "vocab": ["expression", "equation", "variable", "equal", "distributive property", "area model"],
        "misconceptions": [
            "Reading the equal sign as 'the answer comes next.'",
            "Believing a variable always has the same value.",
            "Breaking apart one factor but failing to include every part.",
        ],
        "smart": "Replace the unknown and check that both sides are equal.",
        "tool": "area_grid",
    },
    6: {
        "name": "fractions",
        "theory": "A fraction is a number built from equal parts of the same whole. The denominator names the size of the parts and the numerator counts them. Models and number lines establish magnitude before procedures.",
        "vocab": ["fraction", "numerator", "denominator", "unit fraction", "equivalent", "benchmark"],
        "misconceptions": [
            "Comparing denominators as if larger always means greater.",
            "Using unequal wholes when comparing fractions.",
            "Treating numerator and denominator as unrelated whole numbers.",
        ],
        "smart": "Keep the whole the same and ask how large one part is.",
        "tool": "fraction_tool",
    },
    7: {
        "name": "measurement and data",
        "theory": "Measurement compares an attribute with a chosen unit. Data displays encode quantities through scales, keys, positions, and categories, so units and labels are part of the mathematics.",
        "vocab": ["unit", "estimate", "elapsed time", "scale", "line plot", "volume", "mass"],
        "misconceptions": [
            "Ignoring units or mixing unlike units.",
            "Reading a graph without checking its scale or key.",
            "Treating clock minutes like base-ten digits.",
        ],
        "smart": "Name the unit and estimate before calculating.",
        "tool": "measurement_log",
    },
    8: {
        "name": "geometry, area, and perimeter",
        "theory": "Geometry classifies figures by defining attributes. Area measures covering in square units; perimeter measures distance around in linear units. Visual models should keep these attributes distinct.",
        "vocab": ["attribute", "parallel", "perpendicular", "quadrilateral", "area", "perimeter"],
        "misconceptions": [
            "Classifying shapes only by appearance or orientation.",
            "Confusing area with perimeter.",
            "Using linear units for area or square units for perimeter.",
        ],
        "smart": "Name the attribute being measured before choosing a method.",
        "tool": "geometry_grid",
    },
    9: {
        "name": "problem solving",
        "theory": "Non-routine problems require representation, organized search, assumptions, and justification. The goal is not immediate speed but a defensible chain of reasoning.",
        "vocab": ["represent", "deduce", "assumption", "estimate", "invariant", "verify"],
        "misconceptions": [
            "Operating on numbers before understanding the situation.",
            "Stopping after finding one case when all cases are requested.",
            "Giving an estimate without stating assumptions.",
        ],
        "smart": "Draw a simpler case and organize what you know.",
        "tool": "problem_organizer",
    },
    10: {
        "name": "review and reflection",
        "theory": "Cumulative review strengthens retrieval and connections across topics. Corrections are most useful when students identify the cause of an error, rebuild the idea, and solve a transfer problem.",
        "vocab": ["strategy", "evidence", "revise", "transfer", "accurate", "efficient"],
        "misconceptions": [
            "Treating a score as more important than error analysis.",
            "Copying a correction without explaining the changed idea.",
            "Choosing a strategy from habit rather than problem structure.",
        ],
        "smart": "Choose a strategy, show evidence, and check reasonableness.",
        "tool": "review_tracker",
    },
}

WEEKLY_ASSESSMENTS = {
    1: {
        "tasks": [
            "Write 7,042 in expanded form and in words.",
            "Use <, >, or =: 6,398 ___ 6,389. Explain using place value.",
            "Round 4,672 to the nearest 10 and nearest 100.",
            "Estimate 2,487 + 1,536 to the nearest hundred, then find the exact sum.",
            "Find 7,000 - 2,856 and check with addition.",
            "Create a four-digit number with 5 hundreds and a 0 in the tens place. Show two forms.",
        ],
        "answers": [
            "7,000 + 40 + 2; seven thousand forty-two.",
            "6,398 > 6,389 because the tens digits are 9 and 8 after the greater places match.",
            "4,670; 4,700.",
            "About 4,000; exact sum 4,023.",
            "4,144; check 4,144 + 2,856 = 7,000.",
            "Answers vary; the number must have 5 in the hundreds place and 0 in the tens place.",
        ],
    },
    2: {
        "tasks": [
            "Find 468 + 357. Estimate first.",
            "Find 902 - 476. Show every regrouping.",
            "Use compensation to solve 702 - 398.",
            "A school collected 385 cans Monday and 276 Tuesday. How many in all?",
            "Maya had 900 points, spent 248, then earned 175. How many points now?",
            "Choose the best method for 1,000 - 497 and explain your choice.",
        ],
        "answers": ["825.", "426.", "304.", "661 cans.", "827 points.", "503; valid methods and explanations vary."],
    },
    3: {
        "tasks": [
            "Draw an array for 4 x 6 and label rows and columns.",
            "Write a multiplication equation for 7 groups of 3.",
            "Find 24 / 4 and explain what the quotient means.",
            "Write the four fact-family equations for 5, 8, and 40.",
            "Explain why 3 x 8 and 8 x 3 have the same product.",
            "There are 6 bags with 4 shells in each. How many shells? Show two representations.",
        ],
        "answers": ["Array has 24 objects.", "7 x 3 = 21.", "6; meaning depends on grouping or sharing interpretation.", "5 x 8 = 40, 8 x 5 = 40, 40 / 5 = 8, 40 / 8 = 5.", "Turn-around arrays contain the same total.", "24 shells."],
    },
    4: {
        "tasks": [
            "Find 7 x 8 using a known-fact strategy.",
            "Find 9 x 6 using 10 groups minus 1 group.",
            "Find 72 / 8 by naming the unknown factor.",
            "Find 84 / 7 and check with multiplication.",
            "A coach puts 96 balls into 12 equal bins. How many per bin?",
            "Write two different equations that connect 6, 7, and 42.",
        ],
        "answers": ["56.", "54.", "9.", "12; 12 x 7 = 84.", "8 balls.", "Examples: 6 x 7 = 42 and 42 / 7 = 6."],
    },
    5: {
        "tasks": [
            "Use the distributive property to find 7 x 16.",
            "Solve n + 348 = 900 and check.",
            "Are 6 x 24 and (6 x 20) + (6 x 4) equal? Prove it.",
            "A theater has 8 rows of 24 seats. Then 37 seats are blocked. How many remain?",
            "Write an equation for: 5 equal boxes contain 45 markers.",
            "Make both sides equal: 7 x 9 = 60 + ___.",
        ],
        "answers": ["112.", "n = 552.", "Yes; both equal 144.", "155 seats.", "5 x b = 45 or 45 / 5 = b; b = 9.", "3."],
    },
    6: {
        "tasks": [
            "Shade and label 3/4 of a rectangle.",
            "Place 1/2, 5/4, and 2 on a number line.",
            "Use <, >, or =: 2/3 ___ 2/5. Explain.",
            "Name two fractions equivalent to 1/2.",
            "Order from least to greatest: 3/8, 3/4, 3/6.",
            "Lena ate 3/8 of one same-size pizza and Omar ate 5/8. Who ate more, and by how much?",
        ],
        "answers": ["Three of four equal parts shaded.", "1/2, 5/4, 2 in that order.", "2/3 > 2/5.", "Examples: 2/4 and 4/8.", "3/8, 3/6, 3/4.", "Omar by 2/8, or 1/4, of a pizza."],
    },
    7: {
        "tasks": [
            "A lesson starts at 9:35 and ends at 10:20. How long is it?",
            "Which is a reasonable mass for a watermelon: 4 grams, 4 kilograms, or 40 kilograms?",
            "A jug holds 2 liters. How many 250-milliliter cups fill it?",
            "Measurements are 3 1/4, 3 1/2, 3 1/4, and 3 3/4 inches. Which occurs most often?",
            "A graph scale counts by 5. A bar reaches the fourth tick above zero. What value is shown?",
            "Explain why units and graph scales must be written with an answer.",
        ],
        "answers": ["45 minutes.", "4 kilograms.", "8 cups.", "3 1/4 inches.", "20.", "Explanation should state that units and scales determine the quantity represented."],
    },
    8: {
        "tasks": [
            "Name three defining attributes of a rectangle.",
            "Find the area and perimeter of a 6-by-4 rectangle.",
            "Draw two different rectangles with area 24 square units.",
            "A square has perimeter 28 units. Find its side length and area.",
            "Classify a quadrilateral with two pairs of parallel sides and four right angles.",
            "Explain how two shapes can have the same area but different perimeters.",
        ],
        "answers": ["Four sides, four right angles, and two pairs of parallel opposite sides.", "Area 24 square units; perimeter 20 units.", "Examples: 1 by 24, 2 by 12, 3 by 8, or 4 by 6.", "Side 7 units; area 49 square units.", "Rectangle; a square is also possible if all sides are equal.", "Valid examples and explanations vary."],
    },
    9: {
        "tasks": [
            "A number is even, greater than 30, less than 50, and its digits total 10. Find all possibilities.",
            "How many different two-digit numbers can be made with 2, 5, and 8 without repeating a digit?",
            "Estimate how many piano keys are in 12 classrooms if each has about 1 piano. State an assumption.",
            "Find a pattern rule for 3, 7, 11, 15, ... and give the tenth term.",
            "A farmer has chickens and goats with 12 heads and 34 legs. How many of each?",
            "Choose one answer above and prove that no case was missed.",
        ],
        "answers": ["46.", "6.", "About 1,056 keys if each piano has 88 keys; assumptions may vary.", "Add 4; tenth term 39.", "7 chickens and 5 goats.", "Proofs vary; look for an organized list, equation, or elimination argument."],
    },
}


def clean(text: str) -> str:
    text = text.replace("`", "")
    text = re.sub(r"\[(.*?)\]\([^)]*\)", r"\1", text)
    return text.strip()


def parse_lesson(path: Path) -> dict:
    lines = path.read_text().splitlines()
    title = clean(lines[0].removeprefix("# ").strip())
    sections: dict[str, list[str]] = {}
    current = None
    for line in lines[1:]:
        if line.startswith("## "):
            current = clean(line[3:])
            sections[current] = []
        elif current and line.strip():
            sections[current].append(line.rstrip())

    def prose(name: str, fallback="") -> str:
        vals = sections.get(name, [])
        return clean(" ".join(x.lstrip("- ").strip() for x in vals)) or fallback

    def items(name: str) -> list[str]:
        vals = sections.get(name, [])
        out = []
        for line in vals:
            stripped = line.strip()
            if re.match(r"^\d+\.\s+", stripped):
                out.append(clean(re.sub(r"^\d+\.\s+", "", stripped)))
            elif stripped.startswith("- "):
                out.append(clean(stripped[2:]))
            elif stripped.endswith(":"):
                continue
            elif stripped and not out:
                out.append(clean(stripped))
        return out

    main_key = next((k for k in sections if k in ("Main Lesson", "Administration", "Interview Problems", "Performance Task: Plan a Mini Class Party")), None)
    practice_key = next((k for k in sections if k in ("Practice", "Assessment Sections", "Interview Prompts", "Grade 4 Preview")), None)
    week_match = re.search(r"week-(\d+)", str(path))
    day_match = re.search(r"day-(\d+)", path.name)
    lesson = {
        "path": path,
        "week": int(week_match.group(1)),
        "day": int(day_match.group(1)),
        "title": title.split(":", 1)[-1].strip(),
        "goal": prose("Goal", "Build accurate and flexible mathematical reasoning."),
        "materials": [clean(x) for x in re.split(r",\s*", prose("Materials", "Pencil, paper"))],
        "number_talk": prose("Number Talk", "What do you notice?"),
        "main": items(main_key) if main_key else [],
        "practice": items(practice_key) if practice_key else [],
        "challenge": prose("Challenge", "Create a related problem and explain its solution."),
        "exit": prose("Exit Check", "Explain one important idea from today."),
        "answers": [],
    }
    if lesson["day"] == 5 and lesson["week"] in WEEKLY_ASSESSMENTS:
        assessment = WEEKLY_ASSESSMENTS[lesson["week"]]
        lesson["practice"] = assessment["tasks"]
        lesson["answers"] = assessment["answers"]
    return lesson


def setup_doc(teacher=False) -> Document:
    doc = Document()
    h.style_document(doc)
    section = doc.sections[0]
    if teacher:
        section.top_margin = Inches(0.48)
        section.bottom_margin = Inches(0.48)
        section.left_margin = Inches(0.62)
        section.right_margin = Inches(0.62)
        normal = doc.styles["Normal"]
        normal.font.size = Pt(9.5)
        normal.paragraph_format.space_after = Pt(3)
        normal.paragraph_format.line_spacing = 1.02
        for style_name, size in (("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 10.5)):
            doc.styles[style_name].font.size = Pt(size)
            doc.styles[style_name].paragraph_format.space_before = Pt(7)
            doc.styles[style_name].paragraph_format.space_after = Pt(3)
    return doc


def add_student_title(doc, lesson, page, subtitle):
    h.add_title(doc, f"{lesson['title']} Mission", subtitle, student_fields=True)
    h.add_page_label(doc, f"Page {page} | Week {lesson['week']}, Day {lesson['day']}")


def add_response_box(doc, label, lines=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    h.set_font(r, size=11, bold=True, color=h.NAVY)
    h.add_lines(doc, lines)


def add_tool_page(doc, lesson):
    tool = DOMAIN[lesson["week"]]["tool"]
    h.page_break(doc)
    add_student_title(doc, lesson, 5, "Reusable math tool")
    h.add_kid_goal(doc, "Use this page to draw, model, organize, and check.", DOMAIN[lesson["week"]]["smart"])

    if tool in ("number_line", "measurement_log"):
        doc.add_heading("Open Number Lines", level=1)
        for _ in range(4):
            p = doc.add_paragraph("O________________________________________________________________________O")
            p.paragraph_format.space_after = Pt(18)
            for run in p.runs:
                h.set_font(run, size=12, color=h.BLUE)
        add_response_box(doc, "Labels, units, or notes:", 4)
    elif tool in ("place_grid", "fact_grid"):
        doc.add_heading("Model and Record", level=1)
        table = doc.add_table(rows=6, cols=4)
        h.set_table_geometry(table, [1.8] * 4)
        headers = ("Model", "Equation", "Check", "Explain")
        for c, text in zip(table.rows[0].cells, headers):
            h.shade(c, h.NAVY)
            h.cell_border(c, color=h.NAVY)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            h.set_font(r, size=10, bold=True, color=h.WHITE)
        for row in table.rows[1:]:
            row.height = Inches(0.62)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for c in row.cells:
                h.cell_border(c)
    elif tool in ("array_grid", "area_grid", "geometry_grid"):
        doc.add_heading("Grid Lab", level=1)
        table = doc.add_table(rows=10, cols=12)
        h.set_table_geometry(table, [0.60] * 12)
        for row in table.rows:
            row.height = Inches(0.32)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            for c in row.cells:
                h.cell_border(c, color="B8CBE0", size="5")
        add_response_box(doc, "What does your drawing prove?", 1)
    elif tool == "fraction_tool":
        doc.add_heading("Fraction Strips and Number Lines", level=1)
        for denominator in (2, 3, 4, 6, 8):
            p = doc.add_paragraph(f"Divide this strip into {denominator} equal parts:")
            p.paragraph_format.space_after = Pt(2)
            table = doc.add_table(rows=1, cols=denominator)
            h.set_table_geometry(table, [7.2 / denominator] * denominator)
            for c in table.rows[0].cells:
                h.cell_border(c, color=h.BLUE)
                c.text = " "
        p = doc.add_paragraph("0  ______________________________________________________________  1  __________________  2")
        for run in p.runs:
            h.set_font(run, size=11, color=h.NAVY)
    else:
        doc.add_heading("Problem-Solving Organizer", level=1)
        for label in ("What I know", "What I need to find", "My model or table", "My equations", "My check and explanation"):
            add_response_box(doc, label, 2 if label != "My model or table" else 4)


FINAL_ASSESSMENT_PAGES = (
    (
        "Part A: Place Value and Computation",
        (
            ("Write 8,305 in expanded form and in words.", 2),
            ("Use <, >, or =: 7,049 ___ 7,094. Explain with place value.", 2),
            ("Round 6,751 to the nearest 10 and nearest 100.", 1),
            ("Find 2,768 + 1,957. Estimate first.", 2),
            ("Find 7,000 - 3,486. Check with addition.", 2),
        ),
    ),
    (
        "Part B: Multiplication, Division, and Algebra",
        (
            ("Find 8 x 7 in two different ways.", 2),
            ("Find 96 / 12. Write the related multiplication equation.", 2),
            ("Solve 6 x n = 54 and explain what n means.", 2),
            ("Use the distributive property to find 7 x 23.", 2),
            ("Eight boxes hold 24 pencils each. Three full boxes are given away. How many pencils remain?", 2),
        ),
    ),
    (
        "Part C: Fractions, Measurement, and Data",
        (
            ("Use <, >, or =: 3/4 ___ 5/8. Prove it.", 2),
            ("Complete the equivalence: 2/3 = ___/12.", 1),
            ("Draw a number line from 0 to 2 and locate 5/4.", 2),
            ("A game starts at 1:35 p.m. and ends at 2:20 p.m. How long does it last?", 1),
            ("Choose the reasonable mass of a bicycle: 12 grams, 12 kilograms, or 120 kilograms.", 1),
            ("Data: 2, 3, 3, 4, 4, 4, 5. What value occurs most often? How many values are greater than 3?", 2),
        ),
    ),
    (
        "Part D: Geometry and Challenge",
        (
            ("Find the area and perimeter of a 9-by-4 rectangle. Include units.", 2),
            ("A square has perimeter 32 units. Find its side length and area.", 2),
            ("Name a quadrilateral with two pairs of parallel sides and four right angles. State one extra fact that would make it a square.", 2),
            ("Challenge: A farm has 10 chickens and goats altogether and 28 legs altogether. How many of each? Prove your answer.", 4),
        ),
    ),
)

FINAL_ASSESSMENT_ANSWERS = (
    "8,000 + 300 + 5; eight thousand three hundred five.",
    "7,049 < 7,094; the hundreds match and 4 tens is less than 9 tens.",
    "6,750; 6,800.",
    "Estimate about 4,700 or 4,800; exact sum 4,725.",
    "3,514; 3,514 + 3,486 = 7,000.",
    "56; two valid strategies or representations are required.",
    "8; 8 x 12 = 96.",
    "n = 9; n is the size of each group or the missing factor.",
    "7 x 20 + 7 x 3 = 140 + 21 = 161.",
    "120 pencils.",
    "3/4 > 5/8 because 3/4 = 6/8.",
    "8.",
    "5/4 is one fourth beyond 1.",
    "45 minutes.",
    "12 kilograms.",
    "4 occurs most often; four values are greater than 3.",
    "Area 36 square units; perimeter 26 units.",
    "Side length 8 units; area 64 square units.",
    "Rectangle; all four sides equal would make it a square.",
    "6 chickens and 4 goats; 6 x 2 + 4 x 4 = 28 legs.",
)


def add_assessment_items(doc, items, start=1):
    for idx, (prompt, lines) in enumerate(items, start):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(f"{idx}. ")
        h.set_font(r1, size=10.5, bold=True, color=h.BLUE)
        r2 = p.add_run(prompt)
        h.set_font(r2, size=10.5, bold=True if "Challenge:" in prompt else False)
        h.add_lines(doc, lines)


def build_final_assessment_student(lesson: dict, out_dir: Path):
    doc = setup_doc()
    h.add_footer(doc, "Grade 3 Summer Math | Final Assessment")
    item_number = 1
    for page_num, (heading, items) in enumerate(FINAL_ASSESSMENT_PAGES, 1):
        if page_num > 1:
            h.page_break(doc)
        h.add_title(doc, "Grade 3 Final Assessment", heading, student_fields=True)
        h.add_page_label(doc, f"Assessment Page {page_num} of 5")
        if page_num == 1:
            h.add_instruction(doc, "Work independently. Show models, equations, units, and checks. Circle any item you want to revisit.")
        add_assessment_items(doc, items, item_number)
        item_number += len(items)

    h.page_break(doc)
    h.add_title(doc, "Grade 3 Final Assessment", "Reflection and scratch space", student_fields=True)
    h.add_page_label(doc, "Assessment Page 5 of 5")
    h.add_kid_goal(doc, "Review your reasoning without changing answers unless you find clear evidence.", "A careful check is part of strong mathematics.")
    for prompt in (
        "An answer I checked carefully:",
        "A problem where I used more than one strategy:",
        "The item that required the most thinking:",
        "A topic I feel ready to teach someone else:",
    ):
        add_response_box(doc, prompt, 2)
    doc.add_heading("Extra Work Space", level=1)
    h.add_lines(doc, 7)

    path = out_dir / "week-10-day-02-student-packet.docx"
    doc.save(path)
    return path


def build_final_assessment_teacher(lesson: dict, out_dir: Path):
    doc = setup_doc(teacher=True)
    h.add_footer(doc, "Teacher Manual | Grade 3 Final Assessment")
    h.add_teacher_title(doc, "Final Assessment Administration Guide", "Independent assessment | 60-75 minutes")
    h.add_page_label(doc, "Purpose and Validity")
    h.add_instruction(doc, "Goal: obtain an accurate picture of independent Grade 3 mastery, strategy use, and reasoning.")
    doc.add_heading("Assessment Theory", level=1)
    doc.add_paragraph("A useful assessment separates what the student can do independently from what the student can do with instruction. Standard directions, neutral responses, sufficient time, and written strategy observations make the results more trustworthy.")
    h.add_teacher_note(doc, "Core rule", "Do not teach, hint, confirm answers, or name a method during the assessment. Record questions and misconceptions for the correction lesson.")
    doc.add_heading("What This Assessment Samples", level=1)
    h.add_teacher_bullets(doc, (
        "Place value, rounding, multi-digit addition, and subtraction.",
        "Multiplication and division through 12, properties, and unknowns.",
        "Fractions, elapsed time, measurement reasonableness, and data.",
        "Geometry, area, perimeter, and a non-routine system problem.",
        "Explanation, representation, checking, independence, and persistence.",
    ))

    h.page_break(doc)
    h.add_teacher_title(doc, "Preparation and Exact Directions")
    h.add_page_label(doc, "Before Testing")
    h.add_teacher_bullets(doc, (
        "Print the five-page student assessment.",
        "Provide a pencil, eraser, ruler, and blank scratch paper.",
        "Remove visible lesson posters and answer aids.",
        "Plan one short movement or water break after Page 2.",
        "Keep an observation sheet; do not write on the student's assessment during testing.",
    ))
    h.add_script(
        doc,
        '"This assessment shows what you can do independently. Read each problem carefully and show your reasoning. You may use the scratch paper and ruler. If you are unsure, circle the item, make your best attempt, and continue."',
        purpose="Establishes consistent conditions without increasing anxiety.",
        listen_for="Requests for clarification that concern directions rather than mathematical methods.",
        respond="Repeat or paraphrase a direction neutrally. Say, 'Use your best mathematical judgment,' when the student asks for confirmation.",
    )
    h.add_script(
        doc,
        '"You have plenty of time. After Page 2, we will take a short break. I cannot help with the mathematics today, but I will record anything we should revisit later."',
        purpose="Protects independence while assuring the student that errors will become learning opportunities.",
    )
    h.add_teacher_note(doc, "Permitted clarification", "You may pronounce a word or explain a non-mathematical direction. Do not define vocabulary, suggest an operation, or point to a representation.")

    for key_page in (0, 1):
        h.page_break(doc)
        h.add_teacher_title(doc, f"Answer Key: Items {1 + key_page * 10}-{10 + key_page * 10}")
        h.add_page_label(doc, "Score Accuracy and Evidence Separately")
        for idx, answer in enumerate(FINAL_ASSESSMENT_ANSWERS[key_page * 10:(key_page + 1) * 10], 1 + key_page * 10):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            r1 = p.add_run(f"{idx}. ")
            h.set_font(r1, size=9.5, bold=True, color=h.BLUE)
            r2 = p.add_run(answer)
            h.set_font(r2, size=9.5)

    h.page_break(doc)
    h.add_teacher_title(doc, "Scoring, Diagnosis, and Follow-Up")
    h.add_page_label(doc, "Use Results to Plan Day 4")
    doc.add_heading("Suggested Scoring", level=1)
    h.add_teacher_bullets(doc, (
        "Items 1-19: 2 points each; 1 for a correct answer and 1 for required work, units, model, or explanation.",
        "Item 20: 4 points; 1 for each correct count, 1 for a valid representation, and 1 for proof.",
        "Total: 42 points. Report domain patterns in addition to the total score.",
    ))
    doc.add_heading("Error Codes", level=1)
    for label, text in (
        ("R - Reading", "Misread a condition, symbol, unit, or question."),
        ("C - Concept", "Used a model or relationship that does not match the mathematics."),
        ("P - Procedure", "Understood the structure but made a computation or regrouping error."),
        ("E - Evidence", "Answer may be correct, but required reasoning, labels, or units are missing."),
        ("K - Check", "An unreasonable answer remained after an available check."),
    ):
        h.add_teacher_note(doc, label, text)
    doc.add_heading("Post-Assessment Script", level=1)
    h.add_script(
        doc,
        '"You completed serious mathematical work today. We will not correct it now. Tell me which problem required the most thinking and where you checked your work."',
        purpose="Ends with reflection while preserving the assessment for later analysis.",
        respond="Name one observed habit such as persistence, organization, careful modeling, or checking.",
    )

    path = out_dir / "week-10-day-02-teacher-guide.docx"
    doc.save(path)
    return path


def build_student(lesson: dict, out_dir: Path):
    if lesson["week"] == 10 and lesson["day"] == 2:
        return build_final_assessment_student(lesson, out_dir)
    doc = setup_doc()
    h.add_footer(doc, f"Grade 3 Summer Math | Week {lesson['week']}, Day {lesson['day']}")
    domain = DOMAIN[lesson["week"]]

    add_student_title(doc, lesson, 1, "Warm-up and worked example")
    h.add_kid_goal(doc, lesson["goal"], domain["smart"])
    doc.add_heading("Math Talk", level=1)
    p = doc.add_paragraph(lesson["number_talk"])
    for run in p.runs:
        h.set_font(run, size=12, bold=True, color=h.NAVY)
    add_response_box(doc, "My idea:", 1)
    doc.add_heading("New Words", level=1)
    table = doc.add_table(rows=2, cols=3)
    h.set_table_geometry(table, [2.4] * 3)
    for idx, word in enumerate(domain["vocab"][:3]):
        c = table.cell(0, idx)
        h.shade(c, h.LIGHT_BLUE)
        h.cell_border(c, color=h.BLUE)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(word)
        h.set_font(r, size=10.5, bold=True, color=h.NAVY)
        c2 = table.cell(1, idx)
        h.cell_border(c2)
        c2.text = "My meaning or picture:"
    doc.add_heading("Worked Example", level=1)
    example_steps = lesson["main"][:2] or [lesson["goal"]]
    h.add_instruction(doc, "Do this example with your teacher. Say what changes and what stays the same.")
    for idx, step in enumerate(example_steps, 1):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{idx}. ")
        h.set_font(r1, size=11, bold=True, color=h.BLUE)
        r2 = p.add_run(step)
        h.set_font(r2, size=11)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("What the example shows: __________________________________________")
    h.set_font(r, size=10.5, bold=True, color=h.NAVY)

    h.page_break(doc)
    add_student_title(doc, lesson, 2, "Learn it and show your thinking")
    h.add_kid_goal(doc, "Follow the learning path, then explain the big idea.", "Draw before asking for a rule.")
    for idx, step in enumerate(lesson["main"] or [lesson["goal"]], 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"[ ] Step {idx}: ")
        h.set_font(r1, size=10.5, bold=True, color=h.BLUE)
        r2 = p.add_run(step)
        h.set_font(r2, size=10.5)
    add_response_box(doc, "My model, drawing, or organized work:", 6)
    add_response_box(doc, "The big idea in my own words:", 2)

    h.page_break(doc)
    add_student_title(doc, lesson, 3, "Practice lab")
    h.add_kid_goal(doc, "Use a model, equation, or explanation for every task.", "Check whether your answer is reasonable.")
    practice = lesson["practice"] or ["Create and solve two examples related to today's goal."]
    for idx, item in enumerate(practice, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{idx}. ")
        h.set_font(r1, size=11, bold=True, color=h.BLUE)
        r2 = p.add_run(item)
        h.set_font(r2, size=10.5)
        h.add_lines(doc, 2 if len(practice) <= 5 else 1)
    add_response_box(doc, "Choose one answer and prove it another way:", 4)

    h.page_break(doc)
    add_student_title(doc, lesson, 4, "Challenge and final check")
    h.add_kid_goal(doc, "Try an unfamiliar problem and defend your reasoning.", "A wrong first idea can still lead to a strong solution.")
    doc.add_heading("Challenge Mission", level=1)
    p = doc.add_paragraph(lesson["challenge"])
    for run in p.runs:
        h.set_font(run, size=12, bold=True, color=h.NAVY)
    add_response_box(doc, "My plan:", 3)
    add_response_box(doc, "My solution and proof:", 4)
    doc.add_heading("Final Check", level=1)
    p = doc.add_paragraph(lesson["exit"])
    for run in p.runs:
        h.set_font(run, size=11, bold=True)
    h.add_lines(doc, 1)
    table = doc.add_table(rows=1, cols=3)
    h.set_table_geometry(table, [2.4] * 3)
    for c, text, fill in zip(table.rows[0].cells, ("I can teach it", "I need practice", "I need help"), (h.LIGHT_BLUE, h.LIGHT_GOLD, h.LIGHT_GRAY)):
        h.shade(c, fill)
        h.cell_border(c)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("[ ] " + text)
        h.set_font(r, size=10, bold=True, color=h.NAVY)

    add_tool_page(doc, lesson)
    path = out_dir / f"week-{lesson['week']:02d}-day-{lesson['day']:02d}-student-packet.docx"
    doc.save(path)
    return path


def material_use(material: str) -> str:
    low = material.lower()
    if any(x in low for x in ("block", "counter", "tile", "disk", "money", "strip", "circle")):
        return "Sort into labeled piles. Use first to build the quantity, then keep the model visible while recording symbols."
    if any(x in low for x in ("grid", "paper", "notebook", "chart", "plot", "template")):
        return "Place beside the student. Use for drawings, organized records, and a second representation."
    if any(x in low for x in ("ruler", "clock", "scale", "cup", "thermometer")):
        return "Check the tool and units before the lesson. Demonstrate one accurate reading, then let the student handle it."
    if "card" in low:
        return "Cut and shuffle beforehand. Reveal only the cards needed for the current example."
    return "Prepare before the lesson and introduce it only when it supports the current mathematical idea."


def build_teacher(lesson: dict, out_dir: Path):
    if lesson["week"] == 10 and lesson["day"] == 2:
        return build_final_assessment_teacher(lesson, out_dir)
    doc = setup_doc(teacher=True)
    h.add_footer(doc, f"Teacher Manual | Week {lesson['week']}, Day {lesson['day']}")
    domain = DOMAIN[lesson["week"]]
    h.add_teacher_title(doc, f"Teaching Manual: {lesson['title']}", "Scripted lesson | 60-75 minutes")
    h.add_page_label(doc, "Lesson Overview")
    h.add_instruction(doc, f"Learning goal: {lesson['goal']}")
    doc.add_heading("Success Criteria", level=1)
    h.add_teacher_bullets(doc, (
        "The student represents the idea with a model, drawing, table, or number line.",
        "The student uses the lesson vocabulary accurately.",
        "The student explains why a strategy works and checks reasonableness.",
        f"The student can respond independently to: {lesson['exit']}",
    ))
    doc.add_heading("Teaching Theory", level=1)
    p = doc.add_paragraph(domain["theory"])
    for run in p.runs:
        h.set_font(run, size=10)
    h.add_teacher_note(doc, "Instructional principle", "Move concrete -> pictorial -> abstract. Do not remove the model until the student can explain the symbolic work.")
    doc.add_heading("Lesson Flow", level=1)
    table = doc.add_table(rows=6, cols=2)
    h.set_table_geometry(table, [1.25, 5.95])
    flow = (("5-8 min", "Number talk"), ("10 min", "Vocabulary and worked example"), ("20 min", "Main lesson"), ("15 min", "Practice"), ("10 min", "Challenge"), ("5 min", "Exit check"))
    for i, (time, activity) in enumerate(flow):
        for j, text in enumerate((time, activity)):
            c = table.cell(i, j)
            h.cell_border(c)
            if j == 0:
                h.shade(c, h.LIGHT_BLUE)
            r = c.paragraphs[0].add_run(text)
            h.set_font(r, size=9.5, bold=(j == 0), color=h.NAVY if j == 0 else h.BLACK)

    h.page_break(doc)
    h.add_teacher_title(doc, "Materials: Preparation and Use")
    h.add_page_label(doc, "Before the Student Arrives")
    table = doc.add_table(rows=max(1, len(lesson["materials"])), cols=2)
    h.set_table_geometry(table, [2.2, 5.0])
    for i, material in enumerate(lesson["materials"]):
        for j, text in enumerate((material, material_use(material))):
            c = table.cell(i, j)
            h.cell_border(c)
            if j == 0:
                h.shade(c, h.LIGHT_BLUE)
            r = c.paragraphs[0].add_run(text)
            h.set_font(r, size=9, bold=(j == 0), color=h.NAVY if j == 0 else h.BLACK)
    doc.add_heading("Table Setup", level=1)
    h.add_teacher_bullets(doc, (
        "Center: the main model or tool for today's lesson.",
        "Student side: packet, pencil, and only the materials currently needed.",
        "Teacher side: answer notes and observation record, kept out of student view.",
        "Leave a clear space where the student can point, move pieces, and explain.",
    ))
    doc.add_heading("Vocabulary to Establish", level=1)
    for word in domain["vocab"]:
        p = doc.add_paragraph()
        r1 = p.add_run(word + ": ")
        h.set_font(r1, size=9.5, bold=True, color=h.BLUE)
        r2 = p.add_run("Ask the student for an example or picture before refining the definition.")
        h.set_font(r2, size=9.5)
    h.add_teacher_note(doc, "Material rule", "The student should touch, draw, measure, or organize the mathematics before copying a procedure.")

    h.page_break(doc)
    h.add_teacher_title(doc, "Phase 1: Launch and Number Talk")
    h.add_page_label(doc, "Exact Teacher Language")
    h.add_script(
        doc,
        f'"Today we are learning to {lesson["goal"][0].lower() + lesson["goal"][1:]} Before I show a method, I want to hear how you see the numbers."',
        purpose="Positions the student as a mathematical thinker and makes prior knowledge visible.",
        listen_for="Useful representations, precise vocabulary, and whether the student checks reasonableness.",
        respond="Acknowledge the idea neutrally, then ask for evidence rather than immediately evaluating it.",
    )
    doc.add_heading("Number Talk Prompt", level=1)
    h.add_script(
        doc,
        f'"{lesson["number_talk"]}"',
        purpose="Surfaces the relationships needed in the main lesson.",
        listen_for=f"Any reasoning connected to {domain['name']}, not only a final answer.",
        respond="Ask: 'Can you show that another way?' Record two distinct strategies when possible.",
    )
    doc.add_heading("Prompt Ladder", level=1)
    for label, text, fill in (
        ("First", "What do you notice? What could you try?", h.PALE_BLUE),
        ("Next", "Can you draw, build, or organize the information?", h.PALE_BLUE),
        ("Then", "Which part of today's vocabulary fits your idea?", h.LIGHT_GOLD),
        ("Last resort", "Solve a smaller related case, then return to the original.", h.LIGHT_GRAY),
    ):
        h.add_teacher_note(doc, label, text, fill=fill, border=h.MID_GRAY)

    h.page_break(doc)
    h.add_teacher_title(doc, "Phase 2: Explicit Teaching")
    h.add_page_label(doc, "Model, Question, Connect")
    for idx, step in enumerate(lesson["main"] or [lesson["goal"]], 1):
        doc.add_heading(f"Teaching Step {idx}", level=2)
        h.add_script(
            doc,
            f'"Let us {step[0].lower() + step[1:]} Tell me what I should do first and why."',
            purpose="Keeps the student cognitively active during explicit instruction.",
            listen_for="A connection between the representation and the mathematical symbols.",
            respond="Model only the next move. Pause and return responsibility to the student.",
        )
    h.add_teacher_note(doc, "Representation check", "Before leaving the model, ask the student to point to every part of the matching equation or statement.")

    h.page_break(doc)
    h.add_teacher_title(doc, "Phase 3: Guided and Independent Practice")
    h.add_page_label(doc, "Release Responsibility Gradually")
    doc.add_heading("I Do -> We Do -> You Do", level=1)
    h.add_teacher_bullets(doc, (
        "I do: complete the packet's worked example while narrating the reason for each move.",
        "We do: solve the first practice task together; the student chooses the representation.",
        "You do: the student completes the remaining tasks while explaining selected answers.",
    ))
    for idx, item in enumerate(lesson["practice"] or ["Create and solve a related example."], 1):
        h.add_script(
            doc,
            f'"Practice task {idx}: {item} Before solving, tell me what the task is asking and which tool might help."',
            listen_for="Correct interpretation before computation.",
            respond="Use the prompt ladder. Avoid completing a step the student can perform.",
        )
    doc.add_heading("Common Misconceptions", level=1)
    for misconception in domain["misconceptions"]:
        h.add_teacher_note(doc, "Watch for", misconception, fill=h.LIGHT_GOLD, border=h.GOLD)
    doc.add_heading("Differentiation", level=1)
    h.add_teacher_note(doc, "Secure", "Remove routine repetition and ask for a second method, proof, or generalization.")
    h.add_teacher_note(doc, "Developing", "Keep the visual model visible and reduce the number size while preserving the same structure.", fill=h.LIGHT_GOLD, border=h.GOLD)
    h.add_teacher_note(doc, "Fragile", "Return to a concrete example, use one question at a time, and delay independent symbols.", fill=h.LIGHT_GRAY, border=h.MID_GRAY)

    h.page_break(doc)
    h.add_teacher_title(doc, "Phase 4: Challenge, Assessment, and Feedback")
    h.add_page_label(doc, "Protect Productive Struggle")
    h.add_script(
        doc,
        f'"Challenge: {lesson["challenge"]} You may draw, make a table, test a simpler case, or work backward. I will give you thinking time before offering a hint."',
        purpose="Builds persistence and strategy selection rather than rapid answer production.",
        listen_for="An organized plan, revision after a failed attempt, and a justification.",
        respond="Praise a specific mathematical action: organization, checking, representation, or explanation.",
    )
    doc.add_heading("Exit Check", level=1)
    h.add_script(
        doc,
        f'"Complete this independently: {lesson["exit"]}"',
        purpose="Provides evidence of the central lesson goal without teacher prompting.",
        respond="After completion, ask one clarifying question only if the written reasoning is ambiguous.",
    )
    doc.add_heading("Evidence Guide", level=1)
    table = doc.add_table(rows=3, cols=3)
    h.set_table_geometry(table, [1.15, 2.55, 3.50])
    evidence = (
        ("Secure", "Correct and independently justified", "Proceed; add a variation or ask the student to teach the idea."),
        ("Developing", "Correct with model or prompt", "Proceed with the model available and begin next lesson with retrieval."),
        ("Fragile", "Incorrect structure or unexplained answer", "Reteach with smaller quantities and a concrete representation."),
    )
    for i, row in enumerate(evidence):
        for j, text in enumerate(row):
            c = table.cell(i, j)
            h.cell_border(c)
            if j == 0:
                h.shade(c, h.LIGHT_BLUE if i == 0 else h.LIGHT_GOLD if i == 1 else h.LIGHT_GRAY)
            r = c.paragraphs[0].add_run(text)
            h.set_font(r, size=8.8, bold=(j == 0), color=h.NAVY if j == 0 else h.BLACK)

    h.page_break(doc)
    h.add_teacher_title(doc, "Answer and Observation Guide")
    h.add_page_label(doc, "Use After Independent Work")
    doc.add_heading("Expected Mathematical Evidence", level=1)
    h.add_teacher_bullets(doc, (
        f"Number talk: a defensible response to '{lesson['number_talk']}' with a model or explanation.",
        f"Main lesson: evidence that the student can {lesson['goal'][0].lower() + lesson['goal'][1:]}",
        "Practice: accurate work with units, labels, and a reasonableness check where appropriate.",
        f"Challenge: an organized attempt at '{lesson['challenge']}' and an explanation of the chosen strategy.",
        f"Exit: an independent response to '{lesson['exit']}'",
    ))
    doc.add_heading("Teacher Verification Routine", level=1)
    h.add_teacher_bullets(doc, (
        "Rework every numerical answer before marking it.",
        "Accept a different method when the reasoning is valid.",
        "For open problems, verify that every condition is satisfied.",
        "For measurement, data, and geometry, require correct units and labels.",
        "For explanations, distinguish a correct statement from a demonstrated reason.",
    ))
    if lesson["answers"]:
        doc.add_heading("Assessment Answer Notes", level=1)
        for idx, answer in enumerate(lesson["answers"], 1):
            p = doc.add_paragraph()
            r1 = p.add_run(f"{idx}. ")
            h.set_font(r1, size=9.2, bold=True, color=h.BLUE)
            r2 = p.add_run(answer)
            h.set_font(r2, size=9.2)
    doc.add_heading("Observation Record", level=1)
    for label in ("A strong strategy the student used", "A misconception to revisit", "Vocabulary used precisely", "Material or representation that helped", "Plan for the next lesson"):
        h.add_lines(doc, 2, label + ":")

    path = out_dir / f"week-{lesson['week']:02d}-day-{lesson['day']:02d}-teacher-guide.docx"
    doc.save(path)
    return path


def build_all():
    made = []
    for lesson_path in sorted(ROOT.glob("week-*/day-*.md")):
        lesson = parse_lesson(lesson_path)
        if lesson["week"] == 1 and lesson["day"] == 1:
            continue
        out_dir = lesson_path.parent / "materials" / f"day-{lesson['day']:02d}"
        out_dir.mkdir(parents=True, exist_ok=True)
        made.append(build_student(lesson, out_dir))
        made.append(build_teacher(lesson, out_dir))
    print(f"Generated {len(made)} DOCX files")


if __name__ == "__main__":
    build_all()
